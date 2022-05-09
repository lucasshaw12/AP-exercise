#! python3
# This is a custom made module which has been used in exercise.py

import docx

# 1st variation
# A function to get all paragraphs and add to a list, this is then joined by newline '\n' characters
# def getText(filename):
# 	doc = docx.Document(filename)
# 	fullText = []
# 	for para in doc.paragraphs:
# 		fullText.append(para.text)
# 	return '\n'.join(fullText)

# 2nd variation
# A function to get all paragraphs and add to a list, this function indents all paragraphs and adds a double space between lines
# def getText(filename):
# 	doc = docx.Document(filename)
# 	fullText = []
# 	for para in doc.paragraphs:
# 		fullText.append(' ' + para.text)
# 	return '\n\n'.join(fullText)

