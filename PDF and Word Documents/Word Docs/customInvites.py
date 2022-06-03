#! python3
# Generate a word doc with custom invitations

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create word doc
doc = docx.Document()

# Get names from guests.txt
def getText(filename):
	doc = docx.Document(filename)
	fullText = []
	for para in doc.paragraphs:
		fullText.append(para.text)
	return fullText
guests = list(getText('guest.docx'))

j = 4 # 4 equates to the num of lines within each invitation
# One invitiation per page
for i in range(len(guests)):
	paraObj1 = doc.add_paragraph('It would be the pleasure to have the company of')
	paraObj1.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paraObj2 = doc.add_paragraph([guests[i]]) # This is the name from guests.txt
	paraObj2.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paraObj3 = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
	paraObj3.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paraObj4 = doc.add_paragraph('12 August 2022')
	paraObj4.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paraObj5 = doc.add_paragraph('at 7 o\'clock')
	paraObj5.alignment = WD_ALIGN_PARAGRAPH.CENTER
	doc.paragraphs[j].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
	j += 5 # add 5, so the next 5 lines (which amounts to a full invite), 
	# is skipped before the next page break

	# Create Styles
	doc.paragraphs[0].style = 'Heading 4'
	doc.paragraphs[1].style = 'Normal'
	doc.paragraphs[2].style = 'Subtitle'
	doc.paragraphs[3].style = 'Normal'
	doc.paragraphs[4].style = 'Subtitle'

# Save the document
doc.save('Invitations.docx')