#! python3

import docx

#####################################################
# Reading Word Docs

# doc = docx.Document('demo.docx')
# print(len(doc.paragraphs))
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].text)
# print(len(doc.paragraphs[1].runs))
# print(doc.paragraphs[1].runs[0].text)
# print(doc.paragraphs[1].runs[1].text)
# print(doc.paragraphs[1].runs[2].text)
# print(doc.paragraphs[1].runs[3].text)

#####################################################
# Getting full text from a .docx file. The below can be used as a module within another python file. 
# See readDocx.py in this directory

# def getText(filename):
# 	doc = docx.Document(filename)
# 	fullText = []
# 	for para in doc.paragraphs:
# 		fullText.append(para.text)
# 	return '\n'.join(fullText)

# import readDocx # This module has been made manually
# print(readDocx.getText('demo.docx'))

#####################################################
# Creating word docs with nondefault styles

# doc = docx.Document('demo.docx')
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[0].style)
# doc.paragraphs[0].style = 'Normal'

# print(doc.paragraphs[1].text)
# print((doc.paragraphs[1].runs[0].text, 
# 	doc.paragraphs[1].runs[1].text, 
# 	doc.paragraphs[1].runs[2].text, 
# 	doc.paragraphs[1].runs[3].text, 
# 	doc.paragraphs[1].runs[4].text))

# doc.paragraphs[1].runs[0].style = 'Quote Char'
# doc.paragraphs[1].runs[1].style = 'Quote Char'
# doc.paragraphs[1].runs[2].underline = True
# doc.paragraphs[1].runs[4].underline = True
# doc.save('demoRestyled-CAN-BE-DELETED.docx')

#####################################################
# Writing word documents

# doc = docx.Document()
# doc.add_paragraph('Hello, World!')
# doc.add_paragraph('Hello, World!', 'Title') # The second argument is a string of the objects style
# doc.save('Helloworld.docx')

# paraObj1 = doc.add_paragraph('This is a second paragraph.')
# paraObj2 = doc.add_paragraph('This is another paragraph.')
# paraObj1.add_run('This text is being added to the second paragraph.')
# doc.save('multipleparagraphs-CAN-BE-DELETED.docx')

#####################################################
# Adding headings

# doc = docx.Document()
# doc.add_heading('Header 0', 0)
# doc.add_heading('Header 1', 1)
# doc.add_heading('Header 2', 2)
# doc.add_heading('Header 3', 3)
# doc.add_heading('Header 4', 4)
# doc.save('headings-CAN-BE-DELETED.docx')

#####################################################
# Adding line and page breaks

# doc = docx.Document()
# doc.add_paragraph('This is on the first page.')
# doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
# doc.add_paragraph('This is on the second page.')
# doc.save('twoPage-CAN-BE-DELETED.docx')

#####################################################
# Adding pictures

# doc = docx.Document()
# doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
# doc.add_picture('zophie.png')
# doc.save('addPictures-CAN-BE-DELETED.docx')


















